# app/main.py
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse, PlainTextResponse
from pydantic import BaseModel, Field, field_validator, model_validator
from typing import List, Dict, Union, Optional, Literal
import io, os, re, base64, pathlib, requests
import fitz  # PyMuPDF
import docx
import openpyxl
from fpdf import FPDF
from fpdf.errors import FPDFException
from PIL import Image
from urllib.parse import urlparse

try:
    import markdown
except Exception:
    markdown = None

SERVICE_TITLE = "File Processor Service API"
SERVICE_VERSION = "6.3.4"
SERVICE_DESCRIPTION = "Extrai texto (PDF/DOCX/XLSX/TXT) e gera PDFs dinâmicos com texto e imagens."

MAX_UPLOAD_BYTES = 20 * 1024 * 1024
MAX_IMAGE_BYTES  = 15 * 1024 * 1024
HTTP_TIMEOUT_SECS = 12
USER_AGENT = f"Mozilla/5.0 (FileProcessorService/{SERVICE_VERSION})"
DISALLOWED_HOSTS = {"localhost", "127.0.0.1", "0.0.0.0"}
_MM_PER_PX = 25.4 / 96.0
_BULLET_INDENT_MM = 4.0
MANUAL_PATH = pathlib.Path("Manual-File-Processor-Service.md")

app = FastAPI(title=SERVICE_TITLE, version=SERVICE_VERSION, description=SERVICE_DESCRIPTION)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ALLOW_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

Align = Literal["L", "C", "R"]
BlockType = Literal["heading", "subheading", "paragraph", "bullet_list", "key_value", "spacer", "image"]

class ImageContent(BaseModel):
    src: Optional[str] = None
    base64_data: Optional[str] = None
    width: Optional[float] = None
    height: Optional[float] = None
    align: Align = "C"

    @model_validator(mode="after")
    def _one_source(self):
        if not self.src and not self.base64_data:
            raise ValueError("Informe ao menos 'src' ou 'base64_data' para a imagem.")
        return self

class ContentBlock(BaseModel):
    type: BlockType
    content: Union[str, List[str], Dict[str, str], int, ImageContent]
    style: Optional[Dict[str, Union[List[int], str]]] = None
    line_height: Optional[float] = None
    align: Optional[Align] = None

    @field_validator("style")
    @classmethod
    def _validate_style(cls, v):
        if v and "background_color" in v:
            bg = v["background_color"]
            if (not isinstance(bg, list)) or len(bg) != 3 or any((not isinstance(x, int) or x < 0 or x > 255) for x in bg):
                raise ValueError("background_color deve ser [R,G,B] com valores 0-255.")
        return v

class PDFOptions(BaseModel):
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    margins_mm: Optional[List[float]] = None
    page_numbers: bool = True
    title_align: Align = "C"
    theme_text_color: Optional[List[int]] = None
    allow_remote_images: bool = True

    @field_validator("theme_text_color")
    @classmethod
    def _validate_color(cls, v):
        if v is None:
            return v
        if (not isinstance(v, list)) or len(v) != 3 or any((not isinstance(x, int) or x < 0 or x > 255) for x in v):
            raise ValueError("theme_text_color deve ser [R,G,B] com valores 0-255.")
        return v

class DynamicPDF(BaseModel):
    filename: str
    title: str
    content_blocks: List[ContentBlock]
    options: Optional[PDFOptions] = PDFOptions()

class PDFWithFooter(FPDF):
    def __init__(self, *args, page_numbers: bool = True, **kwargs):
        super().__init__(*args, **kwargs)
        self.page_numbers = page_numbers

    def footer(self):
        if not self.page_numbers:
            return
        self.set_y(-15)
        try:
            self.set_font("DejaVu", "", 8)
        except Exception:
            self.set_font("Arial", "", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Página {self.page_no()}", align="C")

def _safe_text(s: str) -> str:
    if not s:
        return ""
    return s.replace("–", "-").replace("—", "-").replace("\u2028", " ").replace("\u2029", " ")

def _read_pdf(contents: bytes) -> str:
    text = []
    with fitz.open(stream=contents, filetype="pdf") as doc:
        for page in doc:
            t = page.get_text("text")
            if t:
                text.append(t)
    return "\n".join(text).strip()

def _read_docx(contents: bytes) -> str:
    buf = io.BytesIO(contents)
    d = docx.Document(buf)
    parts = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            row_vals = [cell.text for cell in row.cells if cell.text]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts).strip()

def _read_xlsx(contents: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        sh = wb[sheet_name]
        parts.append(f"# {sheet_name}")
        for row in sh.iter_rows(values_only=True):
            row_vals = [str(v) for v in row if v is not None]
            if row_vals:
                parts.append(" ; ".join(row_vals))
    return "\n".join(parts).strip()

def _infer_ext(filename: str) -> str:
    return (os.path.splitext(filename or "")[1] or "").lower()

def _extract_text(filename: str, content_type: str, contents: bytes) -> str:
    ext = _infer_ext(filename)
    ct = (content_type or "").lower()
    try:
        if "pdf" in ct or ext == ".pdf":
            return _read_pdf(contents)
        elif "wordprocessingml" in ct or ext == ".docx":
            return _read_docx(contents)
        elif "spreadsheetml" in ct or ext == ".xlsx":
            return _read_xlsx(contents)
        elif "text/plain" in ct or ext == ".txt":
            return contents.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"Tipo de arquivo não suportado: {content_type or ext}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha na extração: {e}")

def _content_width(pdf: FPDF) -> float:
    return pdf.w - pdf.l_margin - pdf.r_margin

def _setup_fonts(pdf: FPDF) -> str:
    try:
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
        pdf.set_font("DejaVu", "", 12)
        return "DejaVu"
    except FPDFException:
        pdf.set_font("Arial", "", 12)
        return "Arial"

def _pdf_bytes(pdf: FPDF) -> bytes:
    out = pdf.output(dest="S")
    return out if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")

def _apply_margins(pdf: FPDF, margins_mm: Optional[List[float]]):
    if not margins_mm:
        return
    if not (1 <= len(margins_mm) <= 3):
        raise HTTPException(status_code=400, detail="margins_mm deve ter 1, 2 ou 3 itens.")
    left = margins_mm[0]
    top = margins_mm[1] if len(margins_mm) >= 2 else margins_mm[0]
    right = margins_mm[2] if len(margins_mm) == 3 else left
    pdf.set_left_margin(left)
    pdf.set_top_margin(top)
    pdf.set_right_margin(right)

def _remaining_space(pdf: FPDF) -> float:
    return pdf.h - pdf.b_margin - pdf.get_y()

def _ensure_space(pdf: FPDF, need_mm: float):
    if _remaining_space(pdf) < max(need_mm, 1.0):
        pdf.add_page()

def _split_lines(pdf: FPDF, w: float, h: float, txt: str) -> List[str]:
    return pdf.multi_cell(w, h, _safe_text(txt), split_only=True)

def _estimate_text_height(pdf: FPDF, w: float, line_h: float, txt: str) -> float:
    lines = _split_lines(pdf, w, line_h, txt)
    return max(1, len(lines)) * line_h

def _is_disallowed_url(url: str) -> bool:
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return True
        host = (parsed.hostname or "").lower()
        return host in DISALLOWED_HOSTS
    except Exception:
        return True

def _image_dims_mm_from_buf(buf: io.BytesIO) -> tuple[float, float]:
    buf.seek(0)
    with Image.open(buf) as im:
        w_px, h_px = im.size
    return w_px * _MM_PER_PX, h_px * _MM_PER_PX

def _to_png_buffer(buf_in: io.BytesIO) -> io.BytesIO:
    buf_in.seek(0)
    with Image.open(buf_in) as im:
        if im.mode not in ("RGB", "RGBA", "L"):
            im = im.convert("RGBA")
        out = io.BytesIO()
        im.save(out, format="PNG")
        out.seek(0)
        return out

def _decode_b64_loose(data: str) -> bytes:
    m = re.match(r"^data:.*?;base64,", data or "", flags=re.IGNORECASE)
    if m:
        data = data[m.end():]
    data = "".join((data or "").split())
    missing = (-len(data)) % 4
    if missing:
        data += "=" * missing
    try:
        return base64.b64decode(data, validate=False)
    except Exception:
        return base64.urlsafe_b64decode(data)

def _fetch_image_to_buffer(img: ImageContent, allow_remote: bool) -> io.BytesIO:
    buf = io.BytesIO()
    if img.base64_data:
        raw = _decode_b64_loose(img.base64_data)
        if len(raw) > MAX_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="Imagem Base64 excede o limite permitido.")
        buf.write(raw); buf.seek(0)
        try:
            with Image.open(buf) as im:
                im.verify()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Base64 não representa imagem válida: {e}")
        buf.seek(0); return buf

    if not img.src:
        raise HTTPException(status_code=400, detail="Imagem sem src/base64_data.")

    if img.src.startswith("http"):
        if not allow_remote:
            raise HTTPException(status_code=400, detail="Carregamento de imagens remotas desabilitado.")
        if _is_disallowed_url(img.src):
            raise HTTPException(status_code=400, detail="URL de imagem não permitida.")
        headers_req = {"User-Agent": USER_AGENT, "Accept": "image/*,*/*;q=0.8", "Referer": "https://www.google.com/"}
        r = requests.get(img.src, headers=headers_req, timeout=HTTP_TIMEOUT_SECS)
        r.raise_for_status()
        if len(r.content) > MAX_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="Imagem remota excede o limite permitido.")
        buf.write(r.content); buf.seek(0)
        return buf

    if not os.path.exists(img.src):
        raise HTTPException(status_code=400, detail=f"Arquivo de imagem não encontrado: {img.src}")
    if os.path.getsize(img.src) > MAX_IMAGE_BYTES:
        raise HTTPException(status_code=413, detail="Imagem local excede o limite permitido.")
    with open(img.src, "rb") as f:
        buf.write(f.read())
    buf.seek(0); return buf

@app.get("/")
def root():
    return {"status": "File Processor Service is running!", "version": SERVICE_VERSION}

@app.get("/manual.md", response_class=PlainTextResponse, tags=["Docs"])
def manual_raw():
    if not MANUAL_PATH.exists():
        raise HTTPException(status_code=404, detail="manual não encontrado")
    return MANUAL_PATH.read_text(encoding="utf-8")

@app.get("/manual", response_class=HTMLResponse, tags=["Docs"])
def manual_html():
    if not MANUAL_PATH.exists():
        raise HTTPException(status_code=404, detail="manual não encontrado")
    text = MANUAL_PATH.read_text(encoding="utf-8")
    if markdown is None:
        return f"<pre style='white-space:pre-wrap;font-family:ui-monospace,monospace'>{text}</pre>"
    body = markdown.markdown(text, extensions=["fenced_code", "tables", "toc", "codehilite"])
    return f"<!doctype html><html><head><meta charset='utf-8'><title>Manual</title></head><body>{body}</body></html>"

@app.post("/process-file", tags=["File Processing"])
async def process_file(
    file: UploadFile = File(...),
    return_as: Literal["json", "txt"] = Query("json"),
    download: bool = Query(False),
):
    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Arquivo muito grande.")
    text = _extract_text(file.filename, file.content_type, contents)
    if return_as == "json":
        return {"filename": file.filename, "content_type": file.content_type, "length": len(text), "extracted_text": text}
    txt_bytes = text.encode("utf-8", errors="replace")
    disp = "attachment" if download else "inline"
    safe_name = (os.path.splitext(file.filename or 'arquivo')[0] or "arquivo").replace(" ", "_") + ".txt"
    headers_resp = {"Content-Disposition": f'{disp}; filename="{safe_name}"'}
    return StreamingResponse(io.BytesIO(txt_bytes), media_type="text/plain; charset=utf-8", headers=headers_resp)

@app.post("/create-pdf", tags=["PDF Generation"])
async def create_dynamic_pdf(
    doc: DynamicPDF,
    download: bool = Query(True),
):
    try:
        pdf = PDFWithFooter(orientation="P", unit="mm", format="A4",
                            page_numbers=doc.options.page_numbers if doc.options else True)
        if doc.options and doc.options.margins_mm:
            _apply_margins(pdf, doc.options.margins_mm)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        active_font = _setup_fonts(pdf)

        if doc.options:
            if doc.options.author: pdf.set_author(doc.options.author)
            if doc.options.subject: pdf.set_subject(doc.options.subject)
            if doc.options.keywords: pdf.set_keywords(doc.options.keywords)
            if doc.options.theme_text_color:
                r, g, b = doc.options.theme_text_color
                pdf.set_text_color(r, g, b)

        cw = _content_width(pdf)
        pdf.set_font(active_font, "B", 18)
        title_align = (doc.options.title_align if doc.options else "C")
        _ensure_space(pdf, 12)
        pdf.multi_cell(w=cw, h=10, txt=_safe_text(doc.title), align=title_align)
        pdf.ln(6)

        for block in doc.content_blocks:
            fill = False
            if getattr(block, "style", None) and "background_color" in block.style:
                r, g, b = block.style["background_color"]
                pdf.set_fill_color(r, g, b); fill = True
            lh = block.line_height

            if block.type == "heading":
                line_h = lh or 8.5; _ensure_space(pdf, line_h + 4)
                pdf.set_font(active_font, "B", 14)
                pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill)
                y = pdf.get_y(); pdf.set_draw_color(200, 200, 200)
                pdf.line(pdf.l_margin, y, pdf.l_margin + cw, y); pdf.ln(4)

            elif block.type == "subheading":
                line_h = lh or 7.0; _ensure_space(pdf, line_h + 2)
                pdf.set_font(active_font, "B", 11)
                pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill); pdf.ln(2)

            elif block.type == "paragraph":
                line_h = lh or 6.0
                need = _estimate_text_height(pdf, cw, line_h, str(block.content)) + 2
                _ensure_space(pdf, need)
                pdf.set_font(active_font, "", 11)
                align = (block.align or "L")
                pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill, align=align); pdf.ln(2)

            elif block.type == "bullet_list":
                if not isinstance(block.content, list):
                    raise HTTPException(status_code=400, detail="bullet_list requer uma lista de strings.")
                line_h = lh or 6.0; pdf.set_font(active_font, "", 11)
                for item in block.content:
                    txt = _safe_text(str(item))
                    need = _estimate_text_height(pdf, cw - _BULLET_INDENT_MM, line_h, txt) + 1.5
                    _ensure_space(pdf, need)
                    x_start = pdf.get_x()
                    pdf.cell(_BULLET_INDENT_MM, line_h, "•", align="L")
                    pdf.multi_cell(cw - _BULLET_INDENT_MM, line_h, txt, fill=fill)
                    pdf.set_x(x_start)
                pdf.ln(1.5)

            elif block.type == "key_value":
                if not isinstance(block.content, dict):
                    raise HTTPException(status_code=400, detail="key_value requer dict {k:v}.")
                line_h = lh or 7.5
                info_line = " | ".join(f"{str(k).strip().capitalize()}: {str(v).strip()}" for k, v in block.content.items())
                need = _estimate_text_height(pdf, cw, line_h, info_line) + 3
                _ensure_space(pdf, need)
                pdf.set_font(active_font, "", 10)
                pdf.multi_cell(cw, line_h, _safe_text(info_line), align="C", fill=fill); pdf.ln(3)

            elif block.type == "spacer":
                if not isinstance(block.content, int):
                    raise HTTPException(status_code=400, detail="spacer requer inteiro (mm).")
                _ensure_space(pdf, block.content); pdf.ln(block.content)

            elif block.type == "image":
                if not isinstance(block.content, ImageContent):
                    raise HTTPException(status_code=400, detail="image requer ImageContent.")
                img = block.content
                raw_buf = _fetch_image_to_buffer(img, allow_remote=doc.options.allow_remote_images if doc.options else True)
                w_mm, h_mm = _image_dims_mm_from_buf(raw_buf)
                png_buf = _to_png_buffer(raw_buf)
                final_w = img.width; final_h = img.height
                if not final_w and not final_h:
                    final_w = min(cw, w_mm); final_h = h_mm * (final_w / w_mm)
                elif not final_w:
                    final_w = w_mm * (img.height / h_mm); final_h = img.height
                elif not final_h:
                    final_h = h_mm * (img.width / w_mm)
                if final_w > cw:
                    scale = cw / final_w; final_w *= scale; final_h *= scale
                _ensure_space(pdf, final_h + 2)
                x_pos = pdf.l_margin
                if img.align == "C": x_pos += (cw - final_w) / 2
                elif img.align == "R": x_pos += (cw - final_w)
                pdf.image(png_buf, x=x_pos, w=final_w, h=final_h, type="PNG")
                pdf.ln(final_h + 2)

            if fill: pdf.set_fill_color(255, 255, 255)

        pdf_bytes = _pdf_bytes(pdf)
        safe_filename = _safe_text(doc.filename).replace(" ", "_") + ".pdf"
        disp = "attachment" if download else "inline"
        headers_resp = {"Content-Disposition": f'{disp}; filename="{safe_filename}"'}
        return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf", headers=headers_resp)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro geral ao gerar o PDF: {e}")
