from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse, PlainTextResponse
from pydantic import BaseModel, Field, field_validator, model_validator
from typing import List, Dict, Union, Optional, Literal
import io
import os
import re
import base64
import fitz  # PyMuPDF
import docx
import openpyxl
from fpdf import FPDF
from fpdf.errors import FPDFException
import requests
from PIL import Image
from urllib.parse import urlparse
import pathlib

# tente usar Python-Markdown para renderizar /manual como HTML
try:
    import markdown  # pip install markdown
except Exception:
    markdown = None

# ==============================================================================
# Configs & Constantes
# ==============================================================================
SERVICE_TITLE = "File Processor Service API"
SERVICE_VERSION = "6.3.4"
SERVICE_DESCRIPTION = "Extrai texto (PDF/DOCX/XLSX/TXT) e gera PDFs dinâmicos com texto e imagens."

MAX_UPLOAD_BYTES = 20 * 1024 * 1024   # 20 MB para /api/process-file
MAX_IMAGE_BYTES  = 15 * 1024 * 1024   # 15 MB por imagem no PDF
HTTP_TIMEOUT_SECS = 12
USER_AGENT = f"Mozilla/5.0 (FileProcessorService/{SERVICE_VERSION})"

DISALLOWED_HOSTS = {"localhost", "127.0.0.1", "0.0.0.0"}  # SSRF básico

# px -> mm (96 dpi)
_MM_PER_PX = 25.4 / 96.0

# Bullet
_BULLET_INDENT_MM = 4.0

# Caminho do manual em Markdown
MANUAL_PATH = pathlib.Path("Manual-File-Processor-Service.md")

# ==============================================================================
# FastAPI
# ==============================================================================
app = FastAPI(title=SERVICE_TITLE, version=SERVICE_VERSION, description=SERVICE_DESCRIPTION)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ALLOW_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==============================================================================
# Models (Pydantic v2)
# ==============================================================================
Align = Literal["L", "C", "R"]
BlockType = Literal["heading", "subheading", "paragraph", "bullet_list", "key_value", "spacer", "image"]

class ImageContent(BaseModel):
    src: Optional[str] = Field(None, description="URL http(s) ou caminho local")
    base64_data: Optional[str] = Field(None, description="Dados Base64 da imagem (alternativa ao src)")
    width: Optional[float] = Field(None, description="Largura em mm")
    height: Optional[float] = Field(None, description="Altura em mm")
    align: Align = Field("C", description="Alinhamento da imagem (L/C/R)")

    @model_validator(mode="after")
    def _one_source(self):
        if not self.src and not self.base64_data:
            raise ValueError("Informe ao menos 'src' ou 'base64_data' para a imagem.")
        return self

class ContentBlock(BaseModel):
    type: BlockType
    content: Union[str, List[str], Dict[str, str], int, ImageContent]
    style: Optional[Dict[str, Union[List[int], str]]] = None
    line_height: Optional[float] = None
    align: Optional[Align] = None  # para paragraph

    @field_validator("style")
    @classmethod
    def _validate_style(cls, v):
        if v and "background_color" in v:
            bg = v["background_color"]
            if (not isinstance(bg, list)) or len(bg) != 3 or any((not isinstance(x, int) or x < 0 or x > 255) for x in bg):
                raise ValueError("background_color deve ser [R,G,B] com valores 0-255.")
        return v

class PDFOptions(BaseModel):
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    margins_mm: Optional[List[float]] = Field(None, description="[left, top, right] em mm (opcional)")
    page_numbers: bool = Field(True, description="Exibir numeração no rodapé")
    title_align: Align = Field("C", description="Alinhamento do título")
    theme_text_color: Optional[List[int]] = Field(None, description="[R,G,B]")
    allow_remote_images: bool = Field(True, description="Permite carregar imagens http(s)")

    @field_validator("theme_text_color")
    @classmethod
    def _validate_color(cls, v):
        if v is None:
            return v
        if (not isinstance(v, list)) or len(v) != 3 or any((not isinstance(x, int) or x < 0 or x > 255) for x in v):
            raise ValueError("theme_text_color deve ser [R,G,B] com valores 0-255.")
        return v

class DynamicPDF(BaseModel):
    filename: str = Field(..., example="meu_relatorio_final")
    title: str = Field(..., example="Relatório de Atividades")
    content_blocks: List[ContentBlock]
    options: Optional[PDFOptions] = Field(default_factory=PDFOptions)

# ==============================================================================
# PDF class
# ==============================================================================
class PDFWithFooter(FPDF):
    def __init__(self, *args, page_numbers: bool = True, **kwargs):
        super().__init__(*args, **kwargs)
        self.page_numbers = page_numbers

    def footer(self):
        if not self.page_numbers:
            return
        self.set_y(-15)
        try:
            self.set_font("DejaVu", "", 8)
        except Exception:
            self.set_font("Arial", "", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Página {self.page_no()}", align="C")

# ==============================================================================
# Utils
# ==============================================================================
def _safe_text(s: str) -> str:
    if not s:
        return ""
    return s.replace("–", "-").replace("—", "-").replace("\u2028", " ").replace("\u2029", " ")

def _read_pdf(contents: bytes) -> str:
    text = []
    with fitz.open(stream=contents, filetype="pdf") as doc:
        for page in doc:
            t = page.get_text("text")
            if t:
                text.append(t)
    return "\n".join(text).strip()

def _read_docx(contents: bytes) -> str:
    buf = io.BytesIO(contents)
    d = docx.Document(buf)
    parts = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            row_vals = [cell.text for cell in row.cells if cell.text]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts).strip()

def _read_xlsx(contents: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        sh = wb[sheet_name]
        parts.append(f"# {sheet_name}")
        for row in sh.iter_rows(values_only=True):
            row_vals = [str(v) for v in row if v is not None]
            if row_vals:
                parts.append(" ; ".join(row_vals))
    return "\n".join(parts).strip()

def _infer_ext(filename: str) -> str:
    return (os.path.splitext(filename or "")[1] or "").lower()

def _extract_text(filename: str, content_type: str, contents: bytes) -> str:
    ext = _infer_ext(filename)
    ct = (content_type or "").lower()
    try:
        if "pdf" in ct or ext == ".pdf":
            return _read_pdf(contents)
        elif "wordprocessingml" in ct or ext == ".docx":
            return _read_docx(contents)
        elif "spreadsheetml" in ct or ext == ".xlsx":
            return _read_xlsx(contents)
        elif "text/plain" in ct or ext == ".txt":
            return contents.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"Tipo de arquivo não suportado: {content_type or ext}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha na extração: {e}")

def _content_width(pdf: FPDF) -> float:
    return pdf.w - pdf.l_margin - pdf.r_margin

def _setup_fonts(pdf: FPDF) -> str:
    """Tenta usar DejaVu; se ausente, cai para Arial."""
    try:
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
        pdf.set_font("DejaVu", "", 12)
        return "DejaVu"
    except FPDFException:
        pdf.set_font("Arial", "", 12)
        return "Arial"

def _pdf_bytes(pdf: FPDF) -> bytes:
    out = pdf.output(dest="S")
    return out if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")

def _apply_margins(pdf: FPDF, margins_mm: Optional[List[float]]):
    if not margins_mm:
        return
    if not (1 <= len(margins_mm) <= 3):
        raise HTTPException(status_code=400, detail="margins_mm deve ter 1 (todos), 2 (esq/top) ou 3 (esq/top/dir).")
    left = margins_mm[0]
    top = margins_mm[1] if len(margins_mm) >= 2 else margins_mm[0]
    right = margins_mm[2] if len(margins_mm) == 3 else left
    pdf.set_left_margin(left)
    pdf.set_top_margin(top)
    pdf.set_right_margin(right)

def _remaining_space(pdf: FPDF) -> float:
    """Espaço restante até a margem inferior, em mm."""
    return pdf.h - pdf.b_margin - pdf.get_y()

def _ensure_space(pdf: FPDF, need_mm: float):
    """Garante espaço antes de iniciar o bloco; senão, cria nova página."""
    if _remaining_space(pdf) < max(need_mm, 1.0):
        pdf.add_page()

def _split_lines(pdf: FPDF, w: float, h: float, txt: str) -> List[str]:
    """Quebra texto como o multi_cell faria (sem desenhar)."""
    return pdf.multi_cell(w, h, _safe_text(txt), split_only=True)

def _estimate_text_height(pdf: FPDF, w: float, line_h: float, txt: str) -> float:
    lines = _split_lines(pdf, w, line_h, txt)
    return max(1, len(lines)) * line_h

def _is_disallowed_url(url: str) -> bool:
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return True
        host = (parsed.hostname or "").lower()
        return host in DISALLOWED_HOSTS
    except Exception:
        return True

def _image_dims_mm_from_buf(buf: io.BytesIO) -> tuple[float, float]:
    buf.seek(0)
    with Image.open(buf) as im:
        w_px, h_px = im.size
    return w_px * _MM_PER_PX, h_px * _MM_PER_PX

def _to_png_buffer(buf_in: io.BytesIO) -> io.BytesIO:
    buf_in.seek(0)
    with Image.open(buf_in) as im:
        if im.mode not in ("RGB", "RGBA", "L"):
            im = im.convert("RGBA")
        out = io.BytesIO()
        im.save(out, format="PNG")
        out.seek(0)
        return out

def _decode_b64_loose(data: str) -> bytes:
    """
    Decodifica Base64 tolerante:
    - remove prefixo data:*;base64,
    - remove espaços/quebras de linha,
    - corrige padding (=),
    - tenta b64decode e, se falhar, urlsafe_b64decode.
    """
    if not isinstance(data, str):
        raise ValueError("base64 precisa ser string.")
    m = re.match(r"^data:.*?;base64,", data, flags=re.IGNORECASE)
    if m:
        data = data[m.end():]
    data = "".join(data.split())
    missing = (-len(data)) % 4
    if missing:
        data += "=" * missing
    try:
        return base64.b64decode(data, validate=False)
    except Exception:
        try:
            return base64.urlsafe_b64decode(data)
        except Exception as e:
            raise ValueError(f"base64 inválido: {e}")

def _fetch_image_to_buffer(img: ImageContent, allow_remote: bool) -> io.BytesIO:
    buf = io.BytesIO()

    # Base64 primeiro
    if img.base64_data:
        try:
            raw = _decode_b64_loose(img.base64_data)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=f"Base64 inválido para imagem: {e}")
        if len(raw) > MAX_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="Imagem Base64 excede o limite permitido.")
        buf.write(raw)
        buf.seek(0)
        # valida se é imagem
        try:
            with Image.open(buf) as im:
                im.verify()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Base64 não representa uma imagem válida: {e}")
        buf.seek(0)
        return buf

    # src obrigatório se não houver base64
    if not img.src:
        raise HTTPException(status_code=400, detail="Imagem sem src/base64_data.")

    # Remota
    if img.src.startswith("http"):
        if not allow_remote:
            raise HTTPException(status_code=400, detail="Carregamento de imagens remotas desabilitado.")
        if _is_disallowed_url(img.src):
            raise HTTPException(status_code=400, detail="URL de imagem não permitida.")

        referer = "https://www.google.com/"
        host = urlparse(img.src).hostname or ""
        if "wikipedia" in host:
            referer = "https://wikipedia.org/"

        headers_req = {
            "User-Agent": USER_AGENT,
            "Accept": "image/*,*/*;q=0.8",
            "Referer": referer,
        }
        try:
            r = requests.get(img.src, headers=headers_req, timeout=HTTP_TIMEOUT_SECS)
            r.raise_for_status()
        except requests.HTTPError as e:
            status = getattr(e.response, "status_code", None)
            if status == 403:
                raise HTTPException(
                    status_code=400,
                    detail="Imagem remota bloqueada (403). Use base64_data ou hospede a imagem em outro domínio."
                )
            raise HTTPException(status_code=400, detail=f"Falha ao baixar imagem: {e}")
        except requests.RequestException as e:
            raise HTTPException(status_code=400, detail=f"Falha ao baixar imagem: {e}")

        if len(r.content) > MAX_IMAGE_BYTES:
            raise HTTPException(status_code=413, detail="Imagem remota excede o limite permitido.")
        buf.write(r.content)
        buf.seek(0)
        return buf

    # Local
    if not os.path.exists(img.src):
        raise HTTPException(status_code=400, detail=f"Arquivo de imagem não encontrado: {img.src}")
    size = os.path.getsize(img.src)
    if size > MAX_IMAGE_BYTES:
        raise HTTPException(status_code=413, detail="Imagem local excede o limite permitido.")
    with open(img.src, "rb") as f:
        buf.write(f.read())
    buf.seek(0)
    return buf

# ==============================================================================
# Endpoints
# ==============================================================================
@app.get("/")
def root():
    return {"status": "File Processor Service is running!", "version": SERVICE_VERSION}

# ---------- Visualização do Manual ----------
@app.get("/manual.md", response_class=PlainTextResponse, tags=["Docs"])
def manual_raw():
    if not MANUAL_PATH.exists():
        raise HTTPException(status_code=404, detail="manual não encontrado")
    return MANUAL_PATH.read_text(encoding="utf-8")

@app.get("/manual", response_class=HTMLResponse, tags=["Docs"])
def manual_html():
    if not MANUAL_PATH.exists():
        raise HTTPException(status_code=404, detail="manual não encontrado")
    text = MANUAL_PATH.read_text(encoding="utf-8")
    if markdown is None:
        # fallback simples
        return f"<pre style='white-space:pre-wrap;font-family:ui-monospace,monospace'>{text}</pre>"
    body = markdown.markdown(text, extensions=["fenced_code", "tables", "toc", "codehilite"])
    html = f"""
    <!doctype html><html lang="pt-br"><head>
      <meta charset="utf-8" />
      <meta name="viewport" content="width=device-width,initial-scale=1" />
      <title>Manual — File Processor Service</title>
      <style>
        body{{font:16px/1.6 system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif; margin:0;}}
        .wrap{{max-width:900px;margin:40px auto;padding:0 20px}}
        pre,code{{font-family:ui-monospace,Menlo,Consolas,monospace}}
        pre{{background:#f6f8fa;padding:12px;border-radius:8px;overflow:auto}}
        table{{border-collapse:collapse;width:100%}} th,td{{border:1px solid #ddd;padding:8px}}
        h1,h2{{border-bottom:1px solid #eee;padding-bottom:.3em}}
        blockquote{{border-left:4px solid #eee;margin:0;padding:.5em 1em;color:#555}}
        .toc{{background:#fbfbfb;border:1px solid #eee;border-radius:8px;padding:12px;}}
      </style>
    </head><body><div class="wrap">{body}</div></body></html>
    """
    return html

# ---------- File processing ----------
@app.post("/api/process-file", tags=["File Processing"])
async def process_file(
    file: UploadFile = File(...),
    return_as: Literal["json", "txt"] = Query("json"),
    download: bool = Query(False, description="Quando return_as=txt, define attachment vs inline"),
):
    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail=f"Arquivo maior que {MAX_UPLOAD_BYTES // (1024*1024)} MB.")

    text = _extract_text(file.filename, file.content_type, contents)

    if return_as == "json":
        return {
            "filename": file.filename,
            "content_type": file.content_type,
            "length": len(text),
            "extracted_text": text,
        }

    # TXT como streaming
    txt_bytes = text.encode("utf-8", errors="replace")
    disp = "attachment" if download else "inline"
    safe_name = (os.path.splitext(file.filename or 'arquivo')[0] or "arquivo").replace(" ", "_") + ".txt"
    headers_resp = {"Content-Disposition": f'{disp}; filename="{safe_name}"'}
    return StreamingResponse(io.BytesIO(txt_bytes), media_type="text/plain; charset=utf-8", headers=headers_resp)

# ---------- PDF generation ----------
@app.post("/api/create-pdf", tags=["PDF Generation"])
async def create_dynamic_pdf(
    doc: DynamicPDF,
    download: bool = Query(True, description="Se True, força download; se False, abre inline no navegador."),
):
    try:
        pdf = PDFWithFooter(
            orientation="P",
            unit="mm",
            format="A4",
            page_numbers=doc.options.page_numbers if doc.options else True
        )

        # Margens antes de add_page
        if doc.options and doc.options.margins_mm:
            _apply_margins(pdf, doc.options.margins_mm)

        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Fontes
        active_font = _setup_fonts(pdf)

        # Metadados
        if doc.options:
            if doc.options.author:
                pdf.set_author(doc.options.author)
            if doc.options.subject:
                pdf.set_subject(doc.options.subject)
            if doc.options.keywords:
                pdf.set_keywords(doc.options.keywords)
            if doc.options.theme_text_color:
                r, g, b = doc.options.theme_text_color
                pdf.set_text_color(r, g, b)

        cw = _content_width(pdf)

        # Título
        pdf.set_font(active_font, "B", 18)
        title_align = (doc.options.title_align if doc.options else "C")
        _ensure_space(pdf, 12)
        pdf.multi_cell(w=cw, h=10, txt=_safe_text(doc.title), align=title_align)
        pdf.ln(6)

        # Blocos
        for block in doc.content_blocks:
            fill = False
            if block.style and "background_color" in block.style:
                r, g, b = block.style["background_color"]
                pdf.set_fill_color(r, g, b)
                fill = True

            lh = block.line_height

            if block.type in ("heading", "subheading", "paragraph", "bullet_list", "key_value"):
                if block.type == "heading":
                    line_h = lh or 8.5
                    _ensure_space(pdf, line_h + 4)
                    pdf.set_font(active_font, "B", 14)
                    pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill)
                    y = pdf.get_y()
                    pdf.set_draw_color(200, 200, 200)
                    pdf.line(pdf.l_margin, y, pdf.l_margin + cw, y)
                    pdf.ln(4)

                elif block.type == "subheading":
                    line_h = lh or 7.0
                    _ensure_space(pdf, line_h + 2)
                    pdf.set_font(active_font, "B", 11)
                    pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill)
                    pdf.ln(2)

                elif block.type == "paragraph":
                    line_h = lh or 6.0
                    need = _estimate_text_height(pdf, cw, line_h, str(block.content)) + 2
                    _ensure_space(pdf, need)
                    pdf.set_font(active_font, "", 11)
                    align = (block.align or "L")
                    pdf.multi_cell(cw, line_h, _safe_text(str(block.content)), fill=fill, align=align)
                    pdf.ln(2)

                elif block.type == "bullet_list":
                    if not isinstance(block.content, list):
                        raise HTTPException(status_code=400, detail="bullet_list requer uma lista de strings.")
                    line_h = lh or 6.0
                    pdf.set_font(active_font, "", 11)
                    for item in block.content:
                        txt = _safe_text(str(item))
                        need = _estimate_text_height(pdf, cw - _BULLET_INDENT_MM, line_h, f"{txt}") + 1.5
                        _ensure_space(pdf, need)
                        x_start = pdf.get_x()
                        pdf.cell(_BULLET_INDENT_MM, line_h, "•", align="L")
                        pdf.multi_cell(cw - _BULLET_INDENT_MM, line_h, txt, fill=fill)
                        pdf.set_x(x_start)
                    pdf.ln(1.5)

                elif block.type == "key_value":
                    if not isinstance(block.content, dict):
                        raise HTTPException(status_code=400, detail="key_value requer um dicionário {k:v}.")
                    line_h = lh or 7.5
                    info_line = " | ".join(f"{str(k).strip().capitalize()}: {str(v).strip()}" for k, v in block.content.items())
                    need = _estimate_text_height(pdf, cw, line_h, info_line) + 3
                    _ensure_space(pdf, need)
                    pdf.set_font(active_font, "", 10)
                    pdf.multi_cell(cw, line_h, _safe_text(info_line), align="C", fill=fill)
                    pdf.ln(3)

            elif block.type == "spacer":
                if not isinstance(block.content, int):
                    raise HTTPException(status_code=400, detail="spacer requer um inteiro (mm).")
                _ensure_space(pdf, block.content)
                pdf.ln(block.content)

            elif block.type == "image":
                if not isinstance(block.content, ImageContent):
                    raise HTTPException(status_code=400, detail="image requer objeto ImageContent.")
                img = block.content
                raw_buf = _fetch_image_to_buffer(img, allow_remote=doc.options.allow_remote_images if doc.options else True)

                # dimensões mm a partir de px
                w_mm, h_mm = _image_dims_mm_from_buf(raw_buf)
                png_buf = _to_png_buffer(raw_buf)

                # calcula finais mantendo proporção
                final_w = img.width
                final_h = img.height
                if not final_w and not final_h:
                    final_w = min(cw, w_mm)
                    final_h = h_mm * (final_w / w_mm)
                elif not final_w:
                    final_w = w_mm * (img.height / h_mm)
                    final_h = img.height
                elif not final_h:
                    final_h = h_mm * (img.width / w_mm)

                if final_w > cw:
                    scale = cw / final_w
                    final_w *= scale
                    final_h *= scale

                _ensure_space(pdf, final_h + 2)

                x_pos = pdf.l_margin
                if img.align == "C":
                    x_pos += (cw - final_w) / 2
                elif img.align == "R":
                    x_pos += (cw - final_w)

                try:
                    pdf.image(png_buf, x=x_pos, w=final_w, h=final_h, type="PNG")
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Falha ao inserir imagem: {e}")
                pdf.ln(final_h + 2)

            # Reset de fill
            if fill:
                pdf.set_fill_color(255, 255, 255)

        pdf_bytes = _pdf_bytes(pdf)
        safe_filename = _safe_text(doc.filename).replace(" ", "_") + ".pdf"
        disp = "attachment" if download else "inline"
        headers_resp = {"Content-Disposition": f'{disp}; filename="{safe_filename}"'}
        return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf", headers=headers_resp)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro geral ao gerar o PDF: {e}")
